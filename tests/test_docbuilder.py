"""docbuilder — spec §4 (`writer`) and §7. Pure assembly, verified by reopening the file.

The documents built here are inspected through python-docx rather than trusted: "it
saved without raising" says nothing about whether the picture, the table or the Czech
diacritics actually landed.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.shared import Inches

from src.agents.charting import derive_levels, render_chart
from src.agents.docbuilder import (
    LEVEL_COLUMN_COUNT,
    LEVEL_COLUMNS_FALLBACK,
    DocumentSections,
    build_document,
)
from tests.ohlc_fixture import EURUSD

ITEMS = [
    {
        "title": "ČNB ponechala sazby",
        "url": "https://example.com/cnb",
        "published": "2026-08-17",
        "snippet": "…",
    },
    {"title": "Dolar oslabil", "url": "https://example.com/usd", "published": "", "snippet": "…"},
]


def sections(**overrides: Any) -> DocumentSections:
    """Czech on purpose: diacritics are a spec-level requirement, not a nice-to-have."""
    base: dict[str, Any] = {
        "title": "EURUSD — analýza trhu",
        "news_heading": "Zprávy",
        "news_analysis": "ČNB ponechala úrokové sazby bez změny (https://example.com/cnb).",
        "chart_heading": "Graf",
        "chart_analysis": "Cena se drží uprostřed rozpětí, klíčová úroveň leží výše.",
        "levels_heading": "Úrovně",
        "level_columns": ["Typ", "Cena", "Zóna", "Testy", "Vstup"],
        "sources_heading": "Zdroje",
    }
    base.update(overrides)
    return DocumentSections(**base)  # type: ignore[typeddict-item]


@pytest.fixture
def chart(tmp_path: Path) -> dict[str, Any]:
    levels = derive_levels(EURUSD, digits=5)
    png = tmp_path / "EURUSD_H1.png"
    render_chart(EURUSD, levels, "EURUSD", "H1", str(png))
    return {"path": str(png), "symbol": "EURUSD", "timeframe": "H1", "levels": levels}


def headings(document: Any) -> list[str]:
    """Every heading, in order. `add_heading(level=0)` yields style "Title", not
    "Heading 0", so filtering on "Heading" alone silently drops the document title."""
    return [
        p.text
        for p in document.paragraphs
        if p.style.name == "Title" or p.style.name.startswith("Heading")
    ]


def body_text(document: Any) -> str:
    cells = " ".join(c.text for t in document.tables for r in t.rows for c in r.cells)
    return " ".join(p.text for p in document.paragraphs) + " " + cells


# --- per-scope structure, spec §4 table ---------------------------------------


def test_news_scope_has_prose_and_sources_but_no_chart(tmp_path: Path) -> None:
    path = build_document(sections(), scope="news", path=str(tmp_path / "r.docx"), items=ITEMS)
    document = Document(path)

    assert headings(document) == ["EURUSD — analýza trhu", "Zprávy", "Zdroje"]
    assert len(document.inline_shapes) == 0
    assert document.tables == []


def test_chart_scope_has_the_image_and_table_but_no_sources(
    tmp_path: Path, chart: dict[str, Any]
) -> None:
    path = build_document(
        sections(), scope="chart", path=str(tmp_path / "r.docx"), chart=chart, items=ITEMS
    )
    document = Document(path)

    assert headings(document) == ["EURUSD — analýza trhu", "Graf", "Úrovně"]
    assert len(document.inline_shapes) == 1
    assert len(document.tables) == 1
    assert "Zdroje" not in headings(document)  # items supplied, but out of scope


def test_both_scope_has_everything(tmp_path: Path, chart: dict[str, Any]) -> None:
    path = build_document(
        sections(), scope="both", path=str(tmp_path / "r.docx"), chart=chart, items=ITEMS
    )
    document = Document(path)

    assert headings(document) == [
        "EURUSD — analýza trhu",
        "Zprávy",
        "Graf",
        "Úrovně",
        "Zdroje",
    ]
    assert len(document.inline_shapes) == 1
    assert len(document.tables) == 1


def test_scope_alone_decides_sections_not_the_models_wording(
    tmp_path: Path, chart: dict[str, Any]
) -> None:
    """Every field is filled, yet a news-scope document still gets no chart."""
    path = build_document(
        sections(), scope="news", path=str(tmp_path / "r.docx"), chart=chart, items=ITEMS
    )
    assert len(Document(path).inline_shapes) == 0


# --- diacritics and headings, spec §7 ----------------------------------------


def test_czech_diacritics_round_trip(tmp_path: Path, chart: dict[str, Any]) -> None:
    """The trap that motivated dropping ReportLab: `ě ř ď` must survive intact."""
    text = sections(news_analysis="Úroveň odporu se potvrdila; ď ě ř š č ž ý á í é ú ů.")
    path = build_document(text, scope="both", path=str(tmp_path / "r.docx"), chart=chart, items=ITEMS)

    content = body_text(Document(path))
    assert "Úroveň" in content
    assert "ď ě ř š č ž ý á í é ú ů" in content
    assert "■" not in content and "?" not in content


def test_headings_use_real_heading_styles_so_a_toc_works(tmp_path: Path) -> None:
    path = build_document(sections(), scope="news", path=str(tmp_path / "r.docx"), items=ITEMS)
    styles = [
        p.style.name for p in Document(path).paragraphs if p.text == "EURUSD — analýza trhu"
    ]
    assert styles == ["Title"]  # add_heading(level=0)


def test_the_chart_is_embedded_at_six_inches(tmp_path: Path, chart: dict[str, Any]) -> None:
    path = build_document(sections(), scope="chart", path=str(tmp_path / "r.docx"), chart=chart)
    assert Document(path).inline_shapes[0].width == Inches(6.0)


# --- the level table ----------------------------------------------------------


def test_the_table_is_a_real_table_with_a_row_per_level(
    tmp_path: Path, chart: dict[str, Any]
) -> None:
    path = build_document(sections(), scope="chart", path=str(tmp_path / "r.docx"), chart=chart)
    table = Document(path).tables[0]

    assert len(table.columns) == LEVEL_COLUMN_COUNT
    assert len(table.rows) == len(chart["levels"]) + 1  # + header
    assert [c.text for c in table.rows[0].cells] == ["Typ", "Cena", "Zóna", "Testy", "Vstup"]


def test_level_kinds_stay_in_english_as_technical_terms(
    tmp_path: Path, chart: dict[str, Any]
) -> None:
    path = build_document(sections(), scope="chart", path=str(tmp_path / "r.docx"), chart=chart)
    kinds = [r.cells[0].text for r in Document(path).tables[0].rows[1:]]
    assert "resistance" in kinds
    assert "session high" in kinds  # underscore rendered as a space


def test_the_entry_candidate_is_marked_in_the_table(
    tmp_path: Path, chart: dict[str, Any]
) -> None:
    path = build_document(sections(), scope="chart", path=str(tmp_path / "r.docx"), chart=chart)
    entries = [r.cells[4].text for r in Document(path).tables[0].rows[1:]]
    assert "long" in entries and "short" in entries


def test_a_session_extreme_has_no_zone_range(tmp_path: Path, chart: dict[str, Any]) -> None:
    """low == high is a single price, and `1.0802 – 1.0802` reads like a bug."""
    path = build_document(sections(), scope="chart", path=str(tmp_path / "r.docx"), chart=chart)
    for row in Document(path).tables[0].rows[1:]:
        if row.cells[0].text.startswith("session"):
            assert row.cells[2].text == ""


def test_a_real_zone_shows_its_envelope(tmp_path: Path, chart: dict[str, Any]) -> None:
    """A clustered zone spans a range, and the table has to show both edges."""
    wide = {
        **chart,
        "levels": [
            {
                "kind": "resistance",
                "price": 1.085,
                "low": 1.084,
                "high": 1.086,
                "touches": 4,
                "strength": 1.0,
                "entry": "short",
            }
        ],
    }
    path = build_document(sections(), scope="chart", path=str(tmp_path / "r.docx"), chart=wide)

    row = Document(path).tables[0].rows[1]
    assert row.cells[2].text == "1.084 \u2013 1.086"


def test_wrong_column_count_falls_back_to_the_default_labels(
    tmp_path: Path, chart: dict[str, Any]
) -> None:
    """Structure must not depend on the model getting a list length right."""
    broken = sections(level_columns=["Typ", "Cena"])
    path = build_document(broken, scope="chart", path=str(tmp_path / "r.docx"), chart=chart)
    header = [c.text for c in Document(path).tables[0].rows[0].cells]
    assert header == list(LEVEL_COLUMNS_FALLBACK)


# --- sources ------------------------------------------------------------------


def test_sources_are_numbered_and_carry_their_urls(tmp_path: Path) -> None:
    path = build_document(sections(), scope="news", path=str(tmp_path / "r.docx"), items=ITEMS)
    listed = [p.text for p in Document(path).paragraphs if p.style.name == "List Number"]

    assert len(listed) == 2
    assert "https://example.com/cnb" in listed[0]
    assert "(2026-08-17)" in listed[0]
    assert "()" not in listed[1]  # blank published date is omitted, not rendered empty


def test_no_items_means_no_empty_sources_heading(tmp_path: Path) -> None:
    """The Tavily-failure path: the news prose notes the gap; a heading over nothing
    reads like a rendering bug."""
    path = build_document(sections(), scope="news", path=str(tmp_path / "r.docx"), items=[])
    assert "Zdroje" not in headings(Document(path))


def test_an_item_without_a_title_falls_back_to_its_url(tmp_path: Path) -> None:
    path = build_document(
        sections(),
        scope="news",
        path=str(tmp_path / "r.docx"),
        items=[{"url": "https://example.com/x", "published": ""}],
    )
    listed = [p.text for p in Document(path).paragraphs if p.style.name == "List Number"]
    assert "https://example.com/x" in listed[0]


# --- guards -------------------------------------------------------------------


def test_an_invalid_scope_is_rejected(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="scope must be one of"):
        build_document(sections(), scope="everything", path=str(tmp_path / "r.docx"))


def test_a_chart_scope_without_the_artifact_at_all_is_rejected(tmp_path: Path) -> None:
    """`None` is a contract violation the router should have prevented — distinct from the
    trader's failure sentinel, which is present but carries no image."""
    with pytest.raises(ValueError, match="needs a chart artifact"):
        build_document(sections(), scope="chart", path=str(tmp_path / "r.docx"), chart=None)


def test_the_trader_failure_sentinel_degrades_instead_of_raising(tmp_path: Path) -> None:
    """path=None, no levels: the heading and the prose stay, the picture and table go."""
    sentinel = {
        "path": None,
        "symbol": "EURUSD",
        "timeframe": "H1",
        "levels": [],
        "commentary": "",
        "error": "TimeoutError('MT5 did not answer')",
    }
    path = build_document(
        sections(), scope="chart", path=str(tmp_path / "r.docx"), chart=sentinel
    )
    document = Document(path)

    assert headings(document) == ["EURUSD — analýza trhu", "Graf"]
    assert len(document.inline_shapes) == 0
    assert document.tables == []
    assert sections()["chart_analysis"] in body_text(document)


def test_a_chart_path_that_vanished_degrades_the_same_way(tmp_path: Path) -> None:
    """Checked before `add_picture`, which would otherwise fail deep inside the save."""
    path = build_document(
        sections(),
        scope="chart",
        path=str(tmp_path / "r.docx"),
        chart={"path": str(tmp_path / "gone.png"), "levels": []},
    )
    assert len(Document(path).inline_shapes) == 0


def test_no_levels_means_no_levels_heading(tmp_path: Path, chart: dict[str, Any]) -> None:
    """One rule throughout: a section with no content gets no heading. A header-only table
    reads like a rendering bug, exactly as an empty source list would."""
    path = build_document(
        sections(), scope="chart", path=str(tmp_path / "r.docx"), chart={**chart, "levels": []}
    )
    document = Document(path)
    assert document.tables == []
    assert "Úrovně" not in headings(document)


def test_both_scope_degrades_to_news_only_when_the_chart_failed(
    tmp_path: Path,
) -> None:
    """The point of the whole change: a dead MT5 still yields a usable report."""
    sentinel = {"path": None, "levels": [], "error": "boom"}
    path = build_document(
        sections(), scope="both", path=str(tmp_path / "r.docx"), chart=sentinel, items=ITEMS
    )
    document = Document(path)

    assert headings(document) == ["EURUSD — analýza trhu", "Zprávy", "Graf", "Zdroje"]
    assert len(document.inline_shapes) == 0


# --- output -------------------------------------------------------------------


def test_the_parent_directory_is_created(tmp_path: Path) -> None:
    target = tmp_path / "reports" / "nested" / "r.docx"
    assert build_document(sections(), scope="news", path=str(target), items=ITEMS) == str(target)
    assert target.is_file()


def test_the_same_sections_build_the_same_structure(tmp_path: Path) -> None:
    """Pure in the sense that matters: the zip's timestamps differ, the content does not."""
    first = Document(
        build_document(sections(), scope="news", path=str(tmp_path / "a.docx"), items=ITEMS)
    )
    second = Document(
        build_document(sections(), scope="news", path=str(tmp_path / "b.docx"), items=ITEMS)
    )
    assert [(p.style.name, p.text) for p in first.paragraphs] == [
        (p.style.name, p.text) for p in second.paragraphs
    ]
