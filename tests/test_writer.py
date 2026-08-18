"""writer node — spec §4. The model is stubbed; docbuilder runs for real.

Running the real builder is deliberate: the writer's job is to hand docbuilder something
it can actually assemble, and a stubbed builder would hide exactly the mismatches worth
catching.
"""

from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from src.agents import writer
from src.agents.charting import derive_levels, render_chart
from tests.ohlc_fixture import EURUSD

NOW = datetime(2026, 8, 18, 12, 30, 15, tzinfo=timezone.utc)

SECTIONS: dict[str, Any] = {
    "title": "EURUSD — analýza",
    "news_heading": "Zprávy",
    "news_analysis": "Sazby zůstaly (https://example.com/a).",
    "chart_heading": "Graf",
    "chart_analysis": "Cena uprostřed rozpětí.",
    "levels_heading": "Úrovně",
    "level_columns": ["Typ", "Cena", "Zóna", "Testy", "Vstup"],
    "sources_heading": "Zdroje",
}

ANALYTICS: dict[str, Any] = {
    "summary": "Sazby zůstaly (https://example.com/a).",
    "items": [
        {
            "title": "Rates held",
            "url": "https://example.com/a",
            "published": "2026-08-17",
            "snippet": "…",
        }
    ],
}


@pytest.fixture(autouse=True)
def report_dir(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    """Never write into the real outputs/ directory from a test."""
    target = tmp_path / "reports"
    monkeypatch.setattr(writer, "REPORT_DIR", str(target))
    return target


@pytest.fixture
def chart(tmp_path: Path) -> dict[str, Any]:
    levels = derive_levels(EURUSD, digits=5)
    png = tmp_path / "EURUSD_H1.png"
    render_chart(EURUSD, levels, "EURUSD", "H1", str(png))
    return {
        "path": str(png),
        "symbol": "EURUSD",
        "timeframe": "H1",
        "levels": levels,
        "generated_at": "2026-08-18T12:00:00+00:00",
        "commentary": "Price sits mid-range.",
    }


@pytest.fixture
def stub_llm(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(writer, "build_sections", lambda **_: (dict(SECTIONS), []))


def base_state(**overrides: Any) -> dict[str, Any]:
    state: dict[str, Any] = {
        "instrument": "EURUSD",
        "timeframe": "H1",
        "scope": "news",
        "language": "cs",
        "analytics_result": ANALYTICS,
        "chart": None,
    }
    state.update(overrides)
    return state


# --- report path --------------------------------------------------------------


def test_report_path_is_timestamped(report_dir: Path) -> None:
    path = Path(writer.report_path("EURUSD", now=NOW))
    assert path.name == "EURUSD_20260818-123015.docx"
    assert path.parent == report_dir


def test_a_retry_never_reuses_a_filename(report_dir: Path) -> None:
    from datetime import timedelta

    first = writer.report_path("EURUSD", now=NOW)
    second = writer.report_path("EURUSD", now=NOW + timedelta(seconds=1))
    assert first != second


# --- the prompt ---------------------------------------------------------------


def _capture_prompt(monkeypatch: pytest.MonkeyPatch) -> dict[str, str]:
    captured: dict[str, str] = {}

    class FakeStructured:
        def invoke(self, prompt: str) -> dict[str, Any]:
            captured["prompt"] = prompt
            return dict(SECTIONS)

    class FakeModel:
        def with_structured_output(self, schema: Any) -> FakeStructured:
            captured["schema"] = schema.__name__
            return FakeStructured()

    monkeypatch.setattr(writer, "chat_model", lambda *a, **k: FakeModel())
    return captured


def test_the_prompt_carries_language_material_and_the_no_advice_rule(
    monkeypatch: pytest.MonkeyPatch, chart: dict[str, Any]
) -> None:
    captured = _capture_prompt(monkeypatch)

    writer.build_sections(
        instrument="EURUSD",
        timeframe="H1",
        scope="both",
        language="cs",
        analytics_result=ANALYTICS,
        chart=chart,
    )

    prompt = captured["prompt"]
    assert captured["schema"] == "DocumentSections"
    assert "cs" in prompt
    assert "https://example.com/a" in prompt
    assert "1.0852" in prompt  # the computed levels, for the model to describe
    assert "Price sits mid-range." in prompt  # the trader's English note, to be rewritten
    assert "no trading advice" in prompt.lower()
    assert "do not recompute" in prompt.lower()


def test_missing_material_is_labelled_rather_than_left_blank(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """The model has to be able to tell "nothing retrieved" from "field omitted"."""
    captured = _capture_prompt(monkeypatch)

    writer.build_sections(
        instrument="EURUSD",
        timeframe="H1",
        scope="news",
        language="en",
        analytics_result={"summary": "", "items": []},
        chart=None,
    )

    assert "none retrieved" in captured["prompt"]


def test_a_non_dict_model_response_is_rejected(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakeStructured:
        def invoke(self, prompt: str) -> Any:
            return "just a string"

    class FakeModel:
        def with_structured_output(self, schema: Any) -> FakeStructured:
            return FakeStructured()

    monkeypatch.setattr(writer, "chat_model", lambda *a, **k: FakeModel())

    with pytest.raises(ValueError, match="expected the sections dict"):
        writer.build_sections(
            instrument="EURUSD",
            timeframe="H1",
            scope="news",
            language="en",
            analytics_result=ANALYTICS,
            chart=None,
        )


# --- the node, happy paths ----------------------------------------------------


def test_run_produces_the_full_document_contract(stub_llm: None) -> None:
    update = writer.run(base_state())

    document = update["document"]
    assert set(document) == {"path", "format", "title", "scope", "language"}
    assert document["format"] == "docx"
    assert document["title"] == SECTIONS["title"]
    assert document["scope"] == "news"
    assert document["language"] == "cs"
    assert Path(document["path"]).is_file()
    assert "errors" not in update


@pytest.mark.parametrize("scope", ["news", "chart", "both"])
def test_every_scope_builds_a_readable_document(
    scope: str, stub_llm: None, chart: dict[str, Any]
) -> None:
    update = writer.run(base_state(scope=scope, chart=chart))

    document = Document(update["document"]["path"])
    assert document.paragraphs[0].text == SECTIONS["title"]
    assert len(document.inline_shapes) == (0 if scope == "news" else 1)


def test_the_instrument_is_normalised_into_the_filename(stub_llm: None) -> None:
    update = writer.run(base_state(instrument="  eurusd  "))
    assert Path(update["document"]["path"]).name.startswith("EURUSD_")


def test_language_defaults_to_english(stub_llm: None) -> None:
    update = writer.run(base_state(language=None))
    assert update["document"]["language"] == "en"


def test_run_logs_one_ok_entry(stub_llm: None) -> None:
    entry = writer.run(base_state())["agent_log"][0]
    assert entry["agent"] == "writer"
    assert entry["status"] == "ok"
    assert isinstance(entry["duration_s"], float)


def test_run_returns_only_keys_this_agent_owns(stub_llm: None) -> None:
    update = writer.run(base_state())
    assert set(update) <= {"document", "messages", "agent_log", "errors"}


def test_run_does_not_mutate_the_state(stub_llm: None) -> None:
    state = base_state()
    snapshot = json.dumps(state, sort_keys=True, default=str)
    writer.run(state)
    assert json.dumps(state, sort_keys=True, default=str) == snapshot


# --- the messages trace -------------------------------------------------------


def test_the_trace_line_is_a_pointer_not_the_prose(stub_llm: None) -> None:
    """`messages` is checkpointed every superstep; the document is already on disk."""
    update = writer.run(base_state())

    message = update["messages"][0]
    assert SECTIONS["title"] in message.content
    assert update["document"]["path"] in message.content
    assert SECTIONS["news_analysis"] not in message.content


def test_a_failed_run_writes_no_trace_line(stub_llm: None) -> None:
    update = writer.run(base_state(scope=None))
    assert "messages" not in update


# --- the node, failure paths --------------------------------------------------


@pytest.mark.parametrize(
    "override,fragment",
    [
        ({"scope": None}, "scope must be one of"),
        ({"scope": "everything"}, "scope must be one of"),
        ({"instrument": ""}, "no instrument"),
        ({"instrument": "   "}, "no instrument"),
    ],
    ids=["scope-none", "scope-invalid", "instrument-empty", "instrument-blank"],
)
def test_bad_state_is_reported_without_raising(
    override: dict[str, Any], fragment: str, stub_llm: None
) -> None:
    update = writer.run(base_state(**override))
    assert update["document"] is None
    assert any(fragment in e for e in update["errors"])
    assert update["agent_log"][0]["status"] == "error"


def test_a_failed_model_call_yields_no_document(monkeypatch: pytest.MonkeyPatch) -> None:
    def boom(**_: Any) -> Any:
        raise RuntimeError("rate limited")

    monkeypatch.setattr(writer, "build_sections", boom)

    update = writer.run(base_state())

    assert update["document"] is None
    assert any("rate limited" in e for e in update["errors"])


def test_a_chart_scope_without_the_artifact_is_reported_not_raised(stub_llm: None) -> None:
    """Should not happen — the router keeps the writer away until the artifact exists — but
    the node must absorb it rather than kill the run."""
    update = writer.run(base_state(scope="chart", chart=None))
    assert update["document"] is None
    assert any("needs a chart artifact" in e for e in update["errors"])


def test_the_mt5_gap_still_produces_a_document(stub_llm: None) -> None:
    """The counterpart of the Tavily-gap test: a failed chart degrades, it does not block."""
    sentinel = {
        "path": None,
        "symbol": "EURUSD",
        "timeframe": "H1",
        "levels": [],
        "commentary": "",
        "error": "TimeoutError('MT5 did not answer')",
    }
    update = writer.run(base_state(scope="both", chart=sentinel))

    assert update["document"] is not None
    document = Document(update["document"]["path"])
    assert len(document.inline_shapes) == 0
    assert document.tables == []


@pytest.mark.parametrize(
    "chart_value,fragment",
    [
        (None, "no chart was requested"),
        ({"path": None, "levels": [], "error": "boom"}, "COULD NOT BE GENERATED"),
    ],
    ids=["absent", "sentinel"],
)
def test_the_prompt_tells_the_model_the_chart_status(
    chart_value: Any, fragment: str, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The paragraph has to be honest about a missing chart, in the target language — which
    is why this is the model's job and not a fixed string in docbuilder."""
    captured = _capture_prompt(monkeypatch)

    writer.build_sections(
        instrument="EURUSD",
        timeframe="H1",
        scope="both",
        language="cs",
        analytics_result=ANALYTICS,
        chart=chart_value,
    )

    assert fragment in captured["prompt"]


def test_a_successful_chart_is_reported_as_such(
    monkeypatch: pytest.MonkeyPatch, chart: dict[str, Any]
) -> None:
    captured = _capture_prompt(monkeypatch)
    writer.build_sections(
        instrument="EURUSD",
        timeframe="H1",
        scope="chart",
        language="cs",
        analytics_result=None,
        chart=chart,
    )
    assert "generated successfully" in captured["prompt"]


def test_the_tavily_gap_still_produces_a_document(stub_llm: None) -> None:
    """Spec §7: an empty-but-present analytics result must still yield a report."""
    update = writer.run(base_state(analytics_result={"summary": "", "items": []}))

    assert update["document"] is not None
    document = Document(update["document"]["path"])
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert SECTIONS["news_heading"] in headings
    assert SECTIONS["sources_heading"] not in headings  # nothing to list


def test_a_missing_analytics_result_is_tolerated(stub_llm: None, chart: dict[str, Any]) -> None:
    """scope "chart" never runs analytics, so the key is legitimately absent."""
    update = writer.run(base_state(scope="chart", chart=chart, analytics_result=None))
    assert update["document"] is not None


def test_cli_json_dump_is_serialisable(stub_llm: None) -> None:
    json.dumps(writer.run(base_state()), default=str)


# --- the model omitting fields ------------------------------------------------


def test_a_missing_heading_is_filled_and_reported() -> None:
    """A live run came back without `sources_heading`, which reached docbuilder as a
    KeyError; the router retried the writer and the second attempt worked, so the run
    survived — at the cost of a step and eight seconds."""
    partial = {k: v for k, v in SECTIONS.items() if k != "sources_heading"}

    sections, filled = writer.complete_sections(
        partial, instrument="EURUSD", timeframe="H1"
    )

    assert sections["sources_heading"] == writer.HEADING_FALLBACKS["sources_heading"]
    assert filled == ["sources_heading"]


def test_a_blank_heading_counts_as_missing() -> None:
    """An empty Word heading is as broken as an absent one."""
    sections, filled = writer.complete_sections(
        {**SECTIONS, "chart_heading": "   "}, instrument="EURUSD", timeframe="H1"
    )
    assert sections["chart_heading"] == writer.HEADING_FALLBACKS["chart_heading"]
    assert "chart_heading" in filled


def test_a_missing_title_falls_back_to_the_instrument() -> None:
    partial = {k: v for k, v in SECTIONS.items() if k != "title"}
    sections, filled = writer.complete_sections(
        partial, instrument="EURUSD", timeframe="H4"
    )
    assert sections["title"] == "EURUSD H4"
    assert "title" in filled


def test_empty_prose_is_not_treated_as_a_defect() -> None:
    """docbuilder legitimately skips an empty chart paragraph, so this is not a fix."""
    _, filled = writer.complete_sections(
        {**SECTIONS, "chart_analysis": ""}, instrument="EURUSD", timeframe="H1"
    )
    assert filled == []


def test_a_wrong_column_count_is_replaced_and_reported() -> None:
    _, filled = writer.complete_sections(
        {**SECTIONS, "level_columns": ["Typ", "Cena"]}, instrument="EURUSD", timeframe="H1"
    )
    assert filled == ["level_columns"]


def test_a_complete_response_needs_no_fixing() -> None:
    sections, filled = writer.complete_sections(
        dict(SECTIONS), instrument="EURUSD", timeframe="H1"
    )
    assert filled == []
    assert sections == SECTIONS


def test_the_run_ships_the_document_and_flags_the_omission(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """The document is the deliverable; a model that keeps omitting fields is a note, not a
    failure."""
    partial = {k: v for k, v in SECTIONS.items() if k != "sources_heading"}
    monkeypatch.setattr(
        writer,
        "build_sections",
        lambda **k: writer.complete_sections(partial, instrument="EURUSD", timeframe="H1"),
    )

    update = writer.run(base_state())

    assert update["document"] is not None
    assert update["agent_log"][0]["status"] == "partial"
    assert any("omitted sources_heading" in e for e in update["errors"])
