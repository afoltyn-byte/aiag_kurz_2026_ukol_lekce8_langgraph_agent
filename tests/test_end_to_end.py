"""End-to-end runs of the compiled graph — spec §7.

Every external boundary is stubbed (MT5, Tavily, all four models) and nothing else is:
the router, the reducers, the checkpointer, `charting`, `docbuilder` and every node body
run for real. These tests are about what the *assembled* graph does — which nodes execute,
in what order, how often, and how it terminates.

`compile_graph()` supplies an `InMemorySaver`, so `clarify` can suspend and resume.
"""

from __future__ import annotations

from collections import Counter
from pathlib import Path
from typing import Any, Callable

import pytest
from docx import Document
from langchain_core.messages import HumanMessage
from langgraph.errors import GraphRecursionError
from langgraph.types import Command

from src.agent import supervisor
from src.agent.config import MAX_CLARIFICATIONS, MAX_STEPS, RECURSION_LIMIT, resolve_model
from src.agent.graph import compile_graph
from src.agent.state import OutputState
from src.agents import analytics, trader, writer
from tests.ohlc_fixture import EURUSD

ITEM = {
    "title": "ECB holds rates",
    "url": "https://example.com/ecb",
    "published": "2026-08-17",
    "snippet": "unchanged",
}
SUMMARY = "Sazby zůstaly bez změny (https://example.com/ecb)."

SECTIONS: dict[str, Any] = {
    "title": "EURUSD — analýza",
    "news_heading": "Zprávy",
    "news_analysis": SUMMARY,
    "chart_heading": "Graf",
    "chart_analysis": "Cena uprostřed rozpětí.",
    "levels_heading": "Úrovně",
    "level_columns": ["Typ", "Cena", "Zóna", "Testy", "Vstup"],
    "sources_heading": "Zdroje",
}

OK = object()

# Captured at import time, before any test monkeypatches it: the model-override test needs
# the real `build_sections` so the invocation config actually flows through it.
REAL_BUILD_SECTIONS = writer.build_sections


# --------------------------------------------------------------------------
# harness
# --------------------------------------------------------------------------


def resolver(
    instrument: str = "EURUSD",
    scope: str = "both",
    *,
    needs_answer: bool = False,
    propose: str | None = None,
) -> Callable[[Any], dict[str, Any]]:
    """A supervisor decision function that behaves like the real one.

    With `needs_answer`, it reports the request as ambiguous until a human answer shows up
    in `messages` — which is how the real supervisor picks up a clarification.
    """

    def decide(state: Any) -> dict[str, Any]:
        answered = any(
            isinstance(m, HumanMessage) for m in (state.get("messages") or [])
        )
        resolved = (not needs_answer) or answered
        return {
            "instrument": instrument if resolved else None,
            "timeframe": "H1",
            "language": "cs",
            "scope": scope if resolved else None,
            "next_agent": propose or ("analytics" if resolved else "clarify"),
            "reason": "scripted",
        }

    return decide


def never_resolves() -> Callable[[Any], dict[str, Any]]:
    """Every answer is unusable — the request stays ambiguous no matter what."""

    def decide(_: Any) -> dict[str, Any]:
        return {
            "instrument": None,
            "timeframe": "H1",
            "language": "cs",
            "scope": None,
            "next_agent": "clarify",
            "reason": "still ambiguous",
        }

    return decide


def install(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
    *,
    decide: Callable[[Any], dict[str, Any]],
    mt5: Any = OK,
    tavily: Any = OK,
    sections: Any = OK,
) -> Counter:
    """Stub the four boundaries and count the calls. Returns the counter."""
    calls: Counter = Counter()

    monkeypatch.setattr(trader, "CHART_DIR", str(tmp_path / "charts"))
    monkeypatch.setattr(writer, "REPORT_DIR", str(tmp_path / "reports"))

    def _decide(state: Any, config: Any = None) -> dict[str, Any]:
        calls["supervisor"] += 1
        return decide(state)

    def _fetch(*_: Any, **__: Any) -> Any:
        calls["mt5"] += 1
        if isinstance(mt5, BaseException):
            raise mt5
        return list(EURUSD)

    def _search(*_: Any, **__: Any) -> Any:
        calls["tavily"] += 1
        if isinstance(tavily, BaseException):
            raise tavily
        return [dict(ITEM)]

    def _sections(**_: Any) -> tuple[dict[str, Any], list[str]]:
        calls["writer"] += 1
        if isinstance(sections, BaseException):
            raise sections
        return dict(SECTIONS), []

    monkeypatch.setattr(supervisor, "decide", _decide)
    monkeypatch.setattr(trader, "fetch_ohlc", _fetch)
    monkeypatch.setattr(trader, "interpret_zones", lambda *a, **k: "Price sits mid-range.")
    monkeypatch.setattr(analytics, "search_news", _search)
    monkeypatch.setattr(analytics, "summarise", lambda *a, **k: SUMMARY)
    monkeypatch.setattr(writer, "build_sections", _sections)
    return calls


def config_for(thread: str, **extra: Any) -> dict[str, Any]:
    """`thread_id` per request, supplied by the caller — never generated in the graph."""
    return {
        "configurable": {"thread_id": thread, **extra},
        "recursion_limit": RECURSION_LIMIT,
    }


def drive(
    graph: Any, payload: Any, config: dict[str, Any]
) -> tuple[list[str], list[Any]]:
    """Stream one invocation; return the node order and any interrupt payloads."""
    visited: list[str] = []
    interrupts: list[Any] = []
    for chunk in graph.stream(payload, config=config, stream_mode="updates"):
        for key, value in chunk.items():
            if key == "__interrupt__":
                interrupts.extend(i.value for i in value)
            else:
                visited.append(key)
    return visited, interrupts


@pytest.fixture
def graph() -> Any:
    return compile_graph()


# --------------------------------------------------------------------------
# one run per scope — spec §7
# --------------------------------------------------------------------------


def test_scope_news_runs_analytics_only(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    calls = install(monkeypatch, tmp_path, decide=resolver(scope="news"))

    visited, _ = drive(graph, {"request": "co je nového na eurodolaru"}, config_for("news"))
    state = graph.get_state(config_for("news")).values

    assert visited == ["supervisor", "analytics", "supervisor", "writer", "supervisor"]
    assert calls["tavily"] == 1
    assert calls["mt5"] == 0  # out of scope, so the trader must never be reached
    # `.get()`, not `[...]`: a channel no node ever wrote is *absent* from state, not
    # None. That is exactly why the router tests `state.get(key) is None` — it has to treat
    # "never produced" and "produced as None" alike.
    assert state.get("chart") is None
    assert state["document"]["scope"] == "news"

    document = Document(state["document"]["path"])
    assert len(document.inline_shapes) == 0
    assert document.tables == []


def test_scope_chart_runs_the_trader_only(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    calls = install(monkeypatch, tmp_path, decide=resolver(scope="chart"))

    visited, _ = drive(graph, {"request": "jak vypadá graf eurodolaru"}, config_for("chart"))
    state = graph.get_state(config_for("chart")).values

    assert visited == ["supervisor", "trader", "supervisor", "writer", "supervisor"]
    assert calls["mt5"] == 1
    assert calls["tavily"] == 0
    assert state.get("analytics_result") is None
    assert Path(state["chart"]["path"]).is_file()

    document = Document(state["document"]["path"])
    assert len(document.inline_shapes) == 1
    assert len(document.tables) == 1


def test_scope_both_runs_each_agent_once_in_order(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    calls = install(monkeypatch, tmp_path, decide=resolver(scope="both"))

    visited, _ = drive(
        graph, {"request": "zprávy i graf na eurodolaru"}, config_for("both")
    )
    state = graph.get_state(config_for("both")).values

    assert visited == [
        "supervisor",
        "analytics",
        "supervisor",
        "trader",
        "supervisor",
        "writer",
        "supervisor",
    ]
    assert calls["tavily"] == 1
    assert calls["mt5"] == 1
    assert calls["writer"] == 1

    document = Document(state["document"]["path"])
    assert len(document.inline_shapes) == 1
    assert len(document.tables) == 1
    assert [e["agent"] for e in state["agent_log"]] == ["analytics", "trader", "writer"]


def test_the_llms_ordering_choice_is_honoured_within_scope(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """`both` needs each agent; which runs first is the supervisor's judgement call."""
    install(monkeypatch, tmp_path, decide=resolver(scope="both", propose="trader"))

    visited, _ = drive(graph, {"request": "graf i zprávy"}, config_for("order"))

    assert visited[:4] == ["supervisor", "trader", "supervisor", "analytics"]


def test_an_out_of_scope_proposal_is_overridden_end_to_end(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """The supervisor proposes the trader on a news-only request; the router refuses."""
    calls = install(monkeypatch, tmp_path, decide=resolver(scope="news", propose="trader"))

    visited, _ = drive(graph, {"request": "jen zprávy"}, config_for("oos"))

    assert "trader" not in visited
    assert calls["mt5"] == 0


# --------------------------------------------------------------------------
# degradation — spec §7
# --------------------------------------------------------------------------


def test_a_failing_tavily_still_produces_a_document(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    calls = install(
        monkeypatch,
        tmp_path,
        decide=resolver(scope="news"),
        tavily=TimeoutError("Tavily did not answer"),
    )

    drive(graph, {"request": "co je nového na eurodolaru"}, config_for("tavily-down"))
    state = graph.get_state(config_for("tavily-down")).values

    assert state["document"] is not None
    assert state["analytics_result"] == {"summary": "", "items": []}
    assert calls["tavily"] == 1  # empty-but-present, so no retry loop
    assert any("Tavily did not answer" in e for e in state["errors"])

    document = Document(state["document"]["path"])
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert SECTIONS["sources_heading"] not in headings  # nothing to list


def test_a_failing_mt5_degrades_instead_of_looping(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """The sentinel's whole purpose: one attempt, then a report that says the chart is
    missing — not eleven attempts ending at `step_limit` with nothing."""
    calls = install(
        monkeypatch,
        tmp_path,
        decide=resolver(scope="chart"),
        mt5=TimeoutError("MT5 MCP did not answer"),
    )

    visited, _ = drive(graph, {"request": "graf eurodolaru"}, config_for("mt5-down"))
    state = graph.get_state(config_for("mt5-down")).values

    assert calls["mt5"] == 1
    assert visited == ["supervisor", "trader", "supervisor", "writer", "supervisor"]
    assert state["step_count"] < MAX_STEPS
    assert state["chart"]["path"] is None
    assert "MT5 MCP did not answer" in state["chart"]["error"]
    assert state["document"] is not None

    document = Document(state["document"]["path"])
    assert len(document.inline_shapes) == 0
    assert document.tables == []


def test_both_dependencies_down_yields_one_document_reporting_both_gaps(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    install(
        monkeypatch,
        tmp_path,
        decide=resolver(scope="both"),
        mt5=TimeoutError("MT5 down"),
        tavily=TimeoutError("Tavily down"),
    )

    drive(graph, {"request": "zprávy i graf"}, config_for("all-down"))
    state = graph.get_state(config_for("all-down")).values

    assert state["document"] is not None
    assert any("MT5 down" in e for e in state["errors"])
    assert any("Tavily down" in e for e in state["errors"])
    assert [e["status"] for e in state["agent_log"]] == ["error", "error", "ok"]


# --------------------------------------------------------------------------
# clarify: interrupt, resume, and the cap — spec §7
# --------------------------------------------------------------------------


def test_an_ambiguous_request_interrupts_then_completes_on_resume(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    install(monkeypatch, tmp_path, decide=resolver(scope="news", needs_answer=True))
    config = config_for("clarify-ok")

    visited, interrupts = drive(graph, {"request": "nějaká analýza"}, config)

    assert visited == ["supervisor"]
    assert len(interrupts) == 1
    payload = interrupts[0]
    assert payload["missing"] == ["instrument", "scope"]
    assert payload["language"] == "cs"
    assert payload["attempt"] == 1
    assert payload["question"]

    resumed, more = drive(graph, Command(resume={"answer": "EURUSD, jen zprávy"}), config)
    state = graph.get_state(config).values

    assert more == []
    assert "clarify" in resumed and "writer" in resumed
    assert state["clarify_count"] == 1
    assert state["document"] is not None


def test_the_answer_reaches_the_supervisor_through_messages(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """The mechanism behind the resume: clarify records the exchange, the supervisor
    re-parses it on its next visit."""
    install(monkeypatch, tmp_path, decide=resolver(scope="news", needs_answer=True))
    config = config_for("clarify-messages")

    drive(graph, {"request": "nějaká analýza"}, config)
    drive(graph, Command(resume={"answer": "EURUSD, jen zprávy"}), config)

    contents = [m.content for m in graph.get_state(config).values["messages"]]
    assert any("EURUSD, jen zprávy" == c for c in contents)


def test_two_unusable_answers_end_the_run_instead_of_asking_a_third_time(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """Spec §5 rule 3: asking a third time is harassment. The run ends with `errors`
    explaining why, and no document."""
    calls = install(monkeypatch, tmp_path, decide=never_resolves())
    config = config_for("clarify-cap")

    _, first = drive(graph, {"request": "něco"}, config)
    _, second = drive(graph, Command(resume={"answer": "nevím"}), config)
    _, third = drive(graph, Command(resume={"answer": "taky nevím"}), config)

    assert len(first) == 1
    assert len(second) == 1
    assert third == []  # the third question is never asked

    state = graph.get_state(config).values
    assert state["clarify_count"] == MAX_CLARIFICATIONS
    assert state.get("document") is None
    assert any("giving up after" in e for e in state["errors"])
    assert calls["tavily"] == 0 and calls["mt5"] == 0 and calls["writer"] == 0


def test_an_empty_answer_still_consumes_one_clarification(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    install(monkeypatch, tmp_path, decide=never_resolves())
    config = config_for("clarify-empty")

    drive(graph, {"request": "něco"}, config)
    drive(graph, Command(resume={"answer": "   "}), config)

    assert graph.get_state(config).values["clarify_count"] == 1


def test_each_question_reports_its_attempt_number(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    install(monkeypatch, tmp_path, decide=never_resolves())
    config = config_for("clarify-attempts")

    _, first = drive(graph, {"request": "něco"}, config)
    _, second = drive(graph, Command(resume={"answer": "nevím"}), config)

    assert first[0]["attempt"] == 1
    assert second[0]["attempt"] == 2


# --------------------------------------------------------------------------
# termination — spec §7
# --------------------------------------------------------------------------


def test_a_supervisor_that_never_finishes_exits_via_step_limit(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """`RECURSION_LIMIT` is 40 and `MAX_STEPS` is 12, so rule 1 has to fire first. A
    GraphRecursionError here would mean the run died instead of ending, losing the trace.
    """
    install(
        monkeypatch,
        tmp_path,
        decide=resolver(scope="news"),
        sections=RuntimeError("writer keeps failing"),
    )
    config = config_for("runaway")

    try:
        visited, _ = drive(graph, {"request": "zprávy"}, config)
    except GraphRecursionError as exc:  # pragma: no cover - the failure we are excluding
        pytest.fail(f"run died on the recursion limit instead of step_limit: {exc}")

    state = graph.get_state(config).values

    assert state["step_count"] == MAX_STEPS
    assert state["document"] is None
    assert any("step limit" in e for e in state["errors"])
    assert visited.count("supervisor") == MAX_STEPS
    assert len(visited) < RECURSION_LIMIT


def test_a_dead_supervisor_model_terminates_through_the_clarify_cap(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """An unreachable supervisor model writes no `instrument` and no `scope`, so rule 3
    treats the request as unresolved and asks — which is the sane fallback. What matters is
    that it still *ends*: two questions, then rule 3 closes the run with `errors`. It never
    spins, and it never reaches an agent on a state nobody could parse.
    """

    def boom(_: Any) -> dict[str, Any]:
        raise RuntimeError("model unreachable")

    calls = install(monkeypatch, tmp_path, decide=boom)
    config = config_for("dead-model")

    _, first = drive(graph, {"request": "cokoli"}, config)
    _, second = drive(graph, Command(resume={"answer": "EURUSD"}), config)
    _, third = drive(graph, Command(resume={"answer": "EURUSD"}), config)

    assert len(first) == 1 and len(second) == 1
    assert third == []  # rule 3 ends it rather than asking again

    state = graph.get_state(config).values
    assert state.get("document") is None
    assert any("model unreachable" in e for e in state["errors"])
    assert calls["tavily"] == 0 and calls["mt5"] == 0 and calls["writer"] == 0


# --------------------------------------------------------------------------
# state contracts — spec §3 and §6
# --------------------------------------------------------------------------


def test_invoke_returns_exactly_the_declared_output_keys(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    install(monkeypatch, tmp_path, decide=resolver(scope="both"))

    returned = graph.invoke({"request": "zprávy i graf"}, config=config_for("out"))

    assert set(returned) == set(OutputState.__annotations__)


def test_the_append_only_keys_accumulate_across_nodes(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """`errors` and `agent_log` carry `operator.add`; without it each node would clobber
    the previous one's entry and the trace would only ever hold the last."""
    install(
        monkeypatch,
        tmp_path,
        decide=resolver(scope="both"),
        mt5=TimeoutError("MT5 down"),
        tavily=TimeoutError("Tavily down"),
    )

    drive(graph, {"request": "zprávy i graf"}, config_for("reducers"))
    state = graph.get_state(config_for("reducers")).values

    assert len(state["agent_log"]) == 3
    assert len(state["errors"]) >= 2


def test_an_unproduced_artifact_is_absent_from_state_not_none(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """A trap worth pinning down: LangGraph stores only channels that were written, so on a
    news-only run the `chart` key is missing entirely rather than holding None. Router rule 4
    reads `state.get(key) is None`, which is what makes both cases behave the same — an
    `in` check or a `state[key]` lookup there would break or raise.
    """
    install(monkeypatch, tmp_path, decide=resolver(scope="news"))

    drive(graph, {"request": "jen zpr\u00e1vy"}, config_for("absent"))
    state = graph.get_state(config_for("absent")).values

    assert "chart" not in state
    assert state.get("chart") is None
    assert supervisor.route_from_supervisor(state) == "done"  # document exists


def test_the_request_is_never_mutated(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    install(monkeypatch, tmp_path, decide=resolver(scope="news"))
    request = "co je nového na eurodolaru"

    drive(graph, {"request": request}, config_for("immutable"))

    assert graph.get_state(config_for("immutable")).values["request"] == request


def test_two_threads_do_not_share_state(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """One `thread_id` per request. A leak here would mix two users' reports."""
    install(monkeypatch, tmp_path, decide=resolver(scope="news"))

    drive(graph, {"request": "první"}, config_for("t-a"))
    drive(graph, {"request": "druhý"}, config_for("t-b"))

    assert graph.get_state(config_for("t-a")).values["request"] == "první"
    assert graph.get_state(config_for("t-b")).values["request"] == "druhý"


def test_a_per_invocation_model_override_reaches_the_node(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """Spec §6: `configurable.models` must travel from `invoke()` into the node, which
    reads it through `langgraph.config.get_config()` rather than a second parameter."""
    install(monkeypatch, tmp_path, decide=resolver(scope="news"))

    seen: dict[str, Any] = {}

    class FakeStructured:
        def invoke(self, prompt: str) -> dict[str, Any]:
            return dict(SECTIONS)

    class FakeModel:
        def with_structured_output(self, schema: Any) -> FakeStructured:
            return FakeStructured()

    def fake_chat_model(agent: str, config: Any = None) -> FakeModel:
        seen[agent] = config
        return FakeModel()

    # The real build_sections this time, with only the model faked underneath it.
    monkeypatch.setattr(writer, "build_sections", REAL_BUILD_SECTIONS)
    monkeypatch.setattr(writer, "chat_model", fake_chat_model)

    drive(
        graph,
        {"request": "zprávy"},
        config_for("override", models={"writer": "gpt-5.4"}),
    )

    assert "writer" in seen, "the writer never resolved a model"
    assert resolve_model("writer", seen["writer"]) == "gpt-5.4"



def test_the_stubbed_boundaries_are_the_only_fakes(
    graph: Any, monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    """Sanity check on the harness itself: the artifacts these runs produce are real files
    written by the real `charting` and `docbuilder`, not fixtures."""
    install(monkeypatch, tmp_path, decide=resolver(scope="both"))

    drive(graph, {"request": "zprávy i graf"}, config_for("real-files"))
    state = graph.get_state(config_for("real-files")).values

    png = Path(state["chart"]["path"])
    docx = Path(state["document"]["path"])
    assert png.read_bytes().startswith(b"\x89PNG")
    assert docx.read_bytes().startswith(b"PK")  # a real zip container
    assert png.parent == tmp_path / "charts"
    assert docx.parent == tmp_path / "reports"
