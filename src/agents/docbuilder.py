"""Word assembly. Pure, no LLM — spec §4 (`writer`) and CLAUDE.md non-negotiable 6.

Same sections in, same document out. All prose and every heading label arrive as
arguments; this module decides *structure* only, never wording. That is what keeps it
language-agnostic without a translation table: the writer's model already produces its
labels in `state["language"]`.

**python-docx, not a PDF library.** ReportLab's built-in Helvetica renders `ě ř ď` as
black boxes with no exception raised, so a broken Czech report would only surface in a
finished file. Word handles Unicode natively.

Like `charting.py` this module raises rather than degrading; the `writer` node catches
and turns failures into `errors` entries.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Mapping, Sequence, TypedDict

from docx import Document
from docx.shared import Inches

CHART_WIDTH_INCHES = 6.0

# Used when the model returns the wrong number of column labels. Structure must not
# depend on the model getting a list length right.
LEVEL_COLUMNS_FALLBACK = ("Level", "Price", "Zone", "Touches", "Entry")

LEVEL_COLUMN_COUNT = len(LEVEL_COLUMNS_FALLBACK)

VALID_SCOPES = ("news", "chart", "both")


class DocumentSections(TypedDict):
    """Everything the model contributes, already written in the target language."""

    title: str
    news_heading: str
    news_analysis: str
    chart_heading: str
    chart_analysis: str
    levels_heading: str
    level_columns: list[str]
    sources_heading: str


def _zone_cell(level: Mapping[str, Any]) -> str:
    """A single price for a session extreme, a range for a real zone."""
    low, high = level.get("low"), level.get("high")
    if low is None or high is None or low == high:
        return ""
    return f"{low} – {high}"


def _usable_image(chart: Mapping[str, Any] | None) -> str | None:
    """The chart's PNG path, or `None` if there is no file to embed.

    Checked here rather than left to `add_picture`, which would fail deep inside the save
    with a path error instead of letting the document degrade gracefully.
    """
    path = (chart or {}).get("path")
    if not path:
        return None
    return str(path) if Path(str(path)).is_file() else None


def _level_columns(sections: Mapping[str, Any]) -> tuple[str, ...]:
    columns = sections.get("level_columns") or []
    if len(columns) != LEVEL_COLUMN_COUNT:
        return LEVEL_COLUMNS_FALLBACK
    return tuple(str(c) for c in columns)


def _add_level_table(document: Any, levels: Sequence[Mapping[str, Any]], columns: Sequence[str]) -> None:
    """A real Word table, not preformatted text — so it stays selectable and sortable.

    `kind` values (`resistance`, `session_high`) are left in English: they are technical
    terms, and translating an enum turns a machine-readable value into prose.
    """
    table = document.add_table(rows=1, cols=LEVEL_COLUMN_COUNT, style="Table Grid")
    for cell, label in zip(table.rows[0].cells, columns):
        cell.text = label

    for level in levels:
        cells = table.add_row().cells
        cells[0].text = str(level.get("kind", "")).replace("_", " ")
        cells[1].text = str(level.get("price", ""))
        cells[2].text = _zone_cell(level)
        cells[3].text = str(level.get("touches", ""))
        cells[4].text = str(level.get("entry") or "")


def _add_sources(document: Any, items: Sequence[Mapping[str, Any]]) -> None:
    for item in items:
        published = item.get("published")
        suffix = f" ({published})" if published else ""
        document.add_paragraph(
            f"{item.get('title') or item.get('url')}{suffix} — {item.get('url')}",
            style="List Number",
        )


def build_document(
    sections: DocumentSections,
    *,
    scope: str,
    path: str,
    chart: Mapping[str, Any] | None = None,
    items: Sequence[Mapping[str, Any]] = (),
) -> str:
    """Assemble the .docx at `path` and return `path`.

    Which sections appear is decided by `scope` alone, per the table in spec §4 — the
    model cannot add or drop a section by wording its output differently.

    Headings go through `add_heading` so a table of contents works later.
    """
    if scope not in VALID_SCOPES:
        raise ValueError(f"scope must be one of {VALID_SCOPES}, got {scope!r}")

    wants_news = scope in ("news", "both")
    wants_chart = scope in ("chart", "both")

    # A chart-bearing scope needs the artifact to exist, but not necessarily a usable
    # image: `trader` writes a sentinel with `path: None` when the fetch fails, and the
    # run is meant to degrade into a document that reports the gap. `chart is None` is a
    # different thing — a contract violation the router should have prevented.
    if wants_chart and chart is None:
        raise ValueError(f"scope {scope!r} needs a chart artifact; got None")

    chart_image = _usable_image(chart)
    levels = (chart or {}).get("levels") or []

    document = Document()
    document.add_heading(sections["title"], level=0)

    if wants_news:
        document.add_heading(sections["news_heading"], level=1)
        document.add_paragraph(sections["news_analysis"])

    if wants_chart:
        document.add_heading(sections["chart_heading"], level=1)
        if chart_image:
            document.add_picture(chart_image, width=Inches(CHART_WIDTH_INCHES))
        if sections["chart_analysis"]:
            document.add_paragraph(sections["chart_analysis"])

        # One rule throughout: a section with no content gets no heading. A table holding
        # only its header row reads like a rendering bug, exactly as an empty source list
        # would. With no levels there is nothing to tabulate.
        if levels:
            document.add_heading(sections["levels_heading"], level=2)
            _add_level_table(document, levels, _level_columns(sections))

    if wants_news and items:
        document.add_heading(sections["sources_heading"], level=1)
        _add_sources(document, items)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path
